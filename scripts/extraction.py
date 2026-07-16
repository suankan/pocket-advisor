"""Shared text-extraction primitives for binary files.

Used by BOTH ingestion paths — email attachments
(extract_attachments.py) and standalone documents
(ingest_documents.py) — so extraction and OCR-caveat behavior cannot
drift between them. Every function here is pure with respect to the
database: Path in, text out.
"""
import shutil
import subprocess
import tempfile
from pathlib import Path

import config


def run_cmd(args, timeout=120):
    return subprocess.run(args, capture_output=True, timeout=timeout)


def ocr_image(image_path: Path):
    """Returns (text, mean_word_confidence)."""
    import pytesseract
    data = pytesseract.image_to_data(
        str(image_path), lang=config.OCR_LANGS,
        output_type=pytesseract.Output.DICT,
    )
    words, confs = [], []
    for word, conf in zip(data["text"], data["conf"]):
        c = float(conf)
        if word.strip() and c >= 0:
            words.append(word)
            confs.append(c)
    text = pytesseract.image_to_string(str(image_path), lang=config.OCR_LANGS)
    mean_conf = sum(confs) / len(confs) if confs else 0.0
    return text, mean_conf


def extract_pdf(path: Path):
    """Returns (text, method, confidence_or_None).

    Always runs both pdftotext and OCR, then prefers whichever yields
    more non-whitespace characters.  This handles hybrid PDFs where
    most text is selectable but embedded-image portions (non-selectable)
    are invisible to pdftotext — the old short-circuit at 40 chars
    silently dropped those portions.
    """
    r = run_cmd(["pdftotext", "-layout", str(path), "-"])
    pt_text = r.stdout.decode("utf-8", errors="replace") if r.returncode == 0 else ""
    pt_len = len("".join(pt_text.split()))

    ocr_text, ocr_conf = "", None
    with tempfile.TemporaryDirectory() as tmp:
        run_cmd(["pdftoppm", "-r", str(config.PDF_OCR_DPI), "-png",
                 str(path), f"{tmp}/page"], timeout=300)
        pages = sorted(Path(tmp).glob("page*.png"))
        if pages:
            texts, confs = [], []
            for page in pages:
                t, c = ocr_image(page)
                texts.append(t)
                confs.append(c)
            ocr_text = "\n\n".join(texts)
            ocr_conf = sum(confs) / len(confs) if confs else None

    ocr_len = len("".join(ocr_text.split())) if ocr_text else 0

    if ocr_len > pt_len:
        return ocr_text, "ocr_tesseract", ocr_conf
    return pt_text, "native_pdftotext", None


def extract_docx(path: Path):
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def extract_xlsx(path: Path):
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def apply_low_confidence_flag(text, conf, source_path: Path):
    """Returns (text, is_low_conf). Below-threshold OCR gets the warning
    prefix and the source file copied to OCR_REVIEW_DIR for human
    verification — junk is never silently indexed as trustworthy."""
    low_conf = conf is not None and conf < config.OCR_LOW_CONFIDENCE
    if low_conf:
        text = "[LOW-CONFIDENCE OCR — VERIFY AGAINST ORIGINAL IMAGE]\n" + (text or "")
        config.OCR_REVIEW_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_path, config.OCR_REVIEW_DIR / source_path.name)
    return text, low_conf
